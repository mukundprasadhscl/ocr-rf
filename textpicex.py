from PyPDF2 import PdfReader
from docx import Document
import io
from PIL import Image as PILImage
import streamlit as st
import easyocr
import fitz



@st.cache_resource
def get_docx_text(docx_file):
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])


@st.cache_resource
def get_pdf_text(pdf_file):
    """
    Extracts text from a PDF file. If the extracted text length is less than 100,
    performs OCR using EasyOCR for better results.
    """
    # Open the PDF file as a binary file
    with open(pdf_file, 'rb') as file:
        pdf_reader = PdfReader(file)
        text = ""

        # Extract text using PyPDF2
        for page in pdf_reader.pages:
            text += page.extract_text()

        # If the extracted text is insufficient, use EasyOCR
        if len(text.strip()) < 100:
            # Log insufficient text length for debugging
            print("Insufficient text extracted using PyPDF2. Running OCR...")

            # Use EasyOCR to extract text from images in the PDF
            file.seek(0)  # Reset file pointer to the beginning of the file
            reader = easyocr.Reader(['en'])  # Initialize EasyOCR with English language

            images = extract_images_from_pdf(file)  # Convert PDF pages to images
            ocr_text = ""
            for image in images:
                # Join the list of strings into a single string before concatenation
                ocr_text += " " + " ".join(reader.readtext(image, detail=0, paragraph=True))

            return ocr_text.strip()  # Return text extracted via OCR if successful

        return text.strip()


def extract_images_from_pdf(pdf_file):
    """
    Converts PDF pages to images using PyMuPDF (fits library).
    Returns a list of images for OCR processing.
    """
    pdf_document = fitz.open(pdf_file)  # Open the PDF with PyMuPDF
    images = []

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        image = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(image)

    pdf_document.close()
    return images


@st.cache_resource
def extract_images_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    images = []
    for page in pdf_reader.pages:
        if "/XObject" in page["/Resources"]:
            x_objects = page["/Resources"]["/XObject"].get_object()
            for obj in x_objects:
                if x_objects[obj]["/Subtype"] == "/Image":
                    try:
                        image_data = x_objects[obj].get_data()
                        image = PILImage.open(io.BytesIO(image_data))
                        images.append(image)
                    except Exception:
                        # Skip unsupported or corrupted images
                        continue
    return images


@st.cache_resource
def extract_images_from_docx(docx_file):
    doc = Document(docx_file)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image = PILImage.open(io.BytesIO(image_data))
                images.append(image)
            except Exception:
                # Skip unsupported or corrupted images
                continue
    return images