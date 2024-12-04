# Working perfectly- without template- deployed- with ocr
from docx.enum.text import WD_LINE_SPACING
import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import re
from PyPDF2 import PdfReader
import json
import base64
from docx import Document
import io
import os
import tempfile
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from PIL import Image as PILImage
import time
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from reportlab.platypus import HRFlowable
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import easyocr
import fitz  # PyMuPDF

# Load environment variables
load_dotenv()
genai_api_key = os.getenv("GOOGLE_API_KEY")


@st.cache_resource
def get_docx_text(docx_file):
    doc = Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])


@st.cache_resource
def get_pdf_text(pdf_file):
    """
    Extracts text from a PDF file. If the extracted text length is less than 100,
    performs OCR using EasyOCR for better results.
    """
    # Open the PDF file as a binary file
    with open(pdf_file, 'rb') as file:
        pdf_reader = PdfReader(file)
        text = ""

        # Extract text using PyPDF2
        for page in pdf_reader.pages:
            text += page.extract_text()

        # If the extracted text is insufficient, use EasyOCR
        if len(text.strip()) < 100:
            # Log insufficient text length for debugging
            print("Insufficient text extracted using PyPDF2. Running OCR...")

            # Use EasyOCR to extract text from images in the PDF
            file.seek(0)  # Reset file pointer to the beginning of the file
            reader = easyocr.Reader(['en'])  # Initialize EasyOCR with English language

            images = extract_images_from_pdf(file)  # Convert PDF pages to images
            ocr_text = ""
            for image in images:
                # Join the list of strings into a single string before concatenation
                ocr_text += " " + " ".join(reader.readtext(image, detail=0, paragraph=True))

            return ocr_text.strip()  # Return text extracted via OCR if successful

        return text.strip()


def extract_images_from_pdf(pdf_file):
    """
    Converts PDF pages to images using PyMuPDF (fits library).
    Returns a list of images for OCR processing.
    """
    pdf_document = fitz.open(pdf_file)  # Open the PDF with PyMuPDF
    images = []

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        image = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(image)

    pdf_document.close()
    return images


def extract_info_with_gemini_mini(text):
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.01)

    # Combined Prompt 1: Personal Info, Education, and Experience
    combined_personal_experience_prompt = f"""
    You are a highly capable resume analysis assistant. Given a resume text, extract and return the following information in JSON format:
    {{
        "personal_info": {{
            "name": "",
            "email": "",
            "phone": "",
            "address": "",
            "LinkedIn": "",
            "date_of_birth": "",
            "nationality": "",
            "father_name": ""
        }},
        "professional_summary": "",
        "career_objective": "",
        "education": [
            {{
                "degree": "",
                "institution": "",
                "year": "",
                "details": ""
            }}
        ],
        "experience": [
            {{
                "title": "",
                "company": "",
                "duration": "",
                "responsibilities": []
            }}
        ]
    }}

    Resume Text: {text}

    Please follow these guidelines:
    1. Place only valid LinkedIn URLs in "LinkedIn".
    2. If name is not present then put name as "Unknown".
    3. If any fields are absent leave it as blank.
    4. Sometimes responsibilities can also be in one sentence or a one line description.


    Return only the JSON object with the extracted information, maintaining the exact structure shown above.
    """

    # Combined Prompt 2: Projects, Skills, and Additional Sections
    combined_projects_skills_other_prompt = f"""
    You are a highly capable resume analysis assistant. Given a resume text, extract and return the following information in JSON format:
    {{
        "projects": [
            {{
                "name": "",
                "description": "",
                "technologies": ""
            }}
        ],
        "skills": {{
            "technical": {{
                "Programming Languages": [],
                "Scripting Languages": [],
                "Databases": [],
                "Monitoring tools": [],
                "Version controllers": [],
                "Operating systems": [],
                "Cloud": [],
                "Devops": [],
                "IAC": [],
                "Automation Tools": [],
                "Data visualization or Report tools ": [],
                "Project Management Tools": [],
                "Full stack": [],
                "App Development": [],
                "IDEs": [],
                "Markup Languages": [],
                "Machine Learning": [],
                "Others": []
            }},
            "soft": []
        }},
        "courses_and_certifications": [
            {{
                "name": "",
                "issuer": "",
                "year": "",
                "type": ""
            }}
        ],
        "additional_sections": {{
            "achievements": [
                {{
                    "title": "",
                    "description": "",
                    "year": ""
                }}
            ],
            "volunteer_work": [
                {{
                    "organization": "",
                    "role": "",
                    "duration": "",
                    "description": ""
                }}
            ],
            "languages": [
                {{
                    "language": "",
                    "proficiency": ""
                }}
            ],
            "awards": [
                {{
                    "name": "",
                    "issuer": "",
                    "year": ""
                }}
            ],
            "publications": [
                {{
                    "title": "",
                    "authors": "",
                    "publication_venue": "",
                    "year": ""
                }}
            ],
            "professional_memberships": [
                {{
                    "organization": "",
                    "role": "",
                    "year": ""
                }}
            ],
            "interests_and_hobbies": [ 
            {{ 
            "name": "", 
            "description": "" 
            }} 
            ]
        }}
    }}

    Resume Text: {text}

    Please follow these guidelines:
    1. Categorize technical skills appropriately. If uncertain, place them in "Others".
    2. Do not combine the skills data with project or professional experience tech stack.
    3. Extract information for all possible sections.
    4. If any fields are absent, leave them blank.
    5. Do not fabricate information.


    Return only the JSON object with the extracted information, maintaining the exact structure shown above.
    """

    def extract_json_from_response(response):
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                return None
        return None

    responses = {
        'personal_experience': extract_json_from_response(model.predict(combined_personal_experience_prompt)),
        'projects_skills_other': extract_json_from_response(model.predict(combined_projects_skills_other_prompt))
    }

    # Combine all responses
    combined_data = {
        'personal_info': responses['personal_experience'].get('personal_info', {}) if responses[
            'personal_experience'] else {},
        'professional_summary': responses['personal_experience'].get('professional_summary', '') if responses[
            'personal_experience'] else '',
        'career_objective': responses['personal_experience'].get('career_objective', '') if responses[
            'personal_experience'] else '',
        'education': responses['personal_experience'].get('education', []) if responses['personal_experience'] else [],
        'experience': responses['personal_experience'].get('experience', []) if responses[
            'personal_experience'] else [],
        'projects': responses['projects_skills_other'].get('projects', []) if responses[
            'projects_skills_other'] else [],
        'skills': responses['projects_skills_other'].get('skills', {}) if responses['projects_skills_other'] else {},
        'courses_and_certifications': responses['projects_skills_other'].get('courses_and_certifications', []) if
        responses['projects_skills_other'] else [],
        'additional_sections': responses['projects_skills_other'].get('additional_sections', {}) if responses[
            'projects_skills_other'] else {}
    }

    #return combined_data

    # Update expected_structure
    expected_structure = {
        "personal_info": {
            "name": "Unknown",
            "email": "",
            "phone": "",
            "address": "",
            "LinkedIn": "",
            "date_of_birth": "",
            "nationality": "",
            "father_name": ""

        },
        "professional_summary": "",
        "career_objective": "",
        "education": [
            {
                "degree": "",
                "institution": "",
                "year": "",
                "details": ""
            }
        ],
        "experience": [
            {
                "title": "",
                "company": "",
                "duration": "",
                "responsibilities": []
            }
        ],
        "projects": [
            {
                "name": "",
                "description": "",
                "technologies": ""
            }
        ],
        "skills": {
            "technical": [],
            "soft": []
        },
        "courses_and_certifications": [
            {
                "name": "",
                "issuer": "",
                "year": "",
                "type": ""
            }
        ],
        # New section
        "additional_sections": {
            "achievements": [
                {
                    "title": "",
                    "description": "",
                    "year": ""
                }
            ],
            "volunteer_work": [
                {
                    "organization": "",
                    "role": "",
                    "duration": "",
                    "description": ""
                }
            ],
            "languages": [
                {
                    "language": "",
                    "proficiency": ""
                }
            ],
            "awards": [
                {
                    "name": "",
                    "issuer": "",
                    "year": ""
                }
            ],
            "publications": [
                {
                    "title": "",
                    "authors": "",
                    "publication_venue": "",
                    "year": ""
                }
            ],
            "professional_memberships": [
                {
                    "organization": "",
                    "role": "",
                    "year": ""
                }
            ],
            "interests_and_hobbies": [
                {
                    "name": "",
                    "description": ""
                }
            ]
        }
    }

    # Fill in missing subfields with defaults
    def fill_defaults(data, structure):
        if isinstance(data, list) and isinstance(structure, list):
            default_item = structure[0] if structure else {}
            return [fill_defaults(item, default_item) for item in data]
        elif isinstance(data, dict) and isinstance(structure, dict):
            return {
                key: fill_defaults(data.get(key, ""), structure[key])
                for key in structure
            }
        else:
            return data if data else ""

    filled_data = fill_defaults(combined_data, expected_structure)

    # Validate LinkedIn URL
    linkedin_url = filled_data['personal_info'].get('LinkedIn', '')
    if linkedin_url and not re.match(r'https?://(www\.)?linkedin\.com/', linkedin_url):
        filled_data['personal_info'].setdefault('other_links', []).append(linkedin_url)
        filled_data['personal_info']['LinkedIn'] = ""

    return filled_data


def extract_info_with_gemini(text):
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.1)

    # Prompt 1: Personal Info and Education
    personal_info_prompt = f"""
    You are a highly capable resume analysis assistant. Given a resume text, extract and return the following information in JSON format:
    {{
        "personal_info": {{
            "name": "",
            "email": "",
            "phone": "",
            "address": "",
            "LinkedIn": "",
            "date_of_birth": "",
            "nationality": "",
            "father_name": ""

        }},
        "professional_summary": "",
        "career_objective": "",
        "education": [
            {{
                "degree": "",
                "institution": "",
                "year": "",
                "details": ""
            }}
        ]
    }}

    Resume Text: {text}

    Please follow these guidelines:
    1. Place only valid LinkedIn URLs in "LinkedIn".
    2. If name is not present then put name as "Unknown".
    3. If any fields are absent, leave it as blank.


    Return only the JSON object with the extracted information, maintaining the exact structure shown above.
    """

    # Prompt 2: Experience
    experience_prompt = f"""
    You are a highly capable resume analysis assistant. Given a resume text, extract and return the following information in JSON format:
    {{
        "experience": [
            {{
                "title": "",
                "company": "",
                "duration": "",
                "responsibilities": []
            }}
        ]
    }}

    Resume Text: {text}
    Return only the JSON object with the extracted information, maintaining the exact structure shown above.
    """

    # Prompt 3: Projects
    projects_prompt = f"""
    You are a highly capable document analysis assistant. Given a resume text, extract and return the following information in JSON format:
    {{
        "projects": [
            {{
                "name": "",
                "description": "",
                "technologies": ""
            }}
        ]
    }}

    Resume Text: {text}
    Please follow these guidelines:
    1. If any fields are absent, leave it as blank.

    Return only the JSON object with the extracted information, maintaining the exact structure shown above.
    """

    # Prompt 4: Skills and Certifications
    skills_cert_prompt = f"""
    You are a highly capable resume analysis assistant. Given a resume text, extract and group technical skills into appropriate categories in JSON format:
    {{
        "skills": {{
            "technical": {{
                "Programming Languages": [],
                "Scripting Languages": [],
                "Databases": [],
                "Monitoring tools": [],
                "Version controllers": [],
                "Operating systems": [],
                "Cloud": [],
                "Devops": [],
                "IAC": [],
                "Automation Tools": [],
                "Data visualization or Report tools ": [],
                "Project Management Tools": [],
                "Full stack": [],
                "App Development": [],
                "IDEs": [],
                "Markup Languages": [],
                "Machine Learning": [],
                "Others": []
            }},
            "soft": []
        }},
        "courses_and_certifications": [
            {{
                "name": "",
                "issuer": "",
                "year": "",
                "type": ""
            }}
        ]
    }}

    Resume Text: {text}

    Please follow these guidelines:
    1. Categorize technical skills appropriately. If uncertain, place them in "Others".
    2. Do not combine the skills data with project or professional experience tech stack.
    3. If any fields are absent, leave them blank.

    Return only the JSON object with the extracted information, maintaining the exact structure shown above.
    """
    other_sections_prompt = f"""
    You are a highly capable resume analysis assistant. Given a resume text, extract and return additional important sections not covered in previous prompts in JSON format:
    {{
        "additional_sections": {{
            "achievements": [
                {{
                    "title": "",
                    "description": "",
                    "year": ""
                }}
            ],
            "volunteer_work": [
                {{
                    "organization": "",
                    "role": "",
                    "duration": "",
                    "description": ""
                }}
            ],
            "languages": [
                {{
                    "language": "",
                    "proficiency": ""
                }}
            ],
            "awards": [
                {{
                    "name": "",
                    "issuer": "",
                    "year": ""
                }}
            ],
            "publications": [
                {{
                    "title": "",
                    "authors": "",
                    "publication_venue": "",
                    "year": ""
                }}
            ],
            "professional_memberships": [
                {{
                    "organization": "",
                    "role": "",
                    "year": ""
                }}
            ],
            "interests_and_hobbies": [ 
            {{ 
            "name": "", 
            "description": "" 
            }} 
            ]
        }}
    }}

    Resume Text: {text}

    Please follow these guidelines:
    1. Extract information for sections that are present in the resume.
    2. If any fields are not present, leave them blank.
    3. Do not fabricate information.

    Return only the JSON object with the extracted information, maintaining the exact structure shown above.
    """

    def extract_json_from_response(response):
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                return None
        return None

    responses = {
        'personal': extract_json_from_response(model.predict(personal_info_prompt)),
        'experience': extract_json_from_response(model.predict(experience_prompt)),
        'projects': extract_json_from_response(model.predict(projects_prompt)),
        'skills_cert': extract_json_from_response(model.predict(skills_cert_prompt)),
        'other_sections': extract_json_from_response(model.predict(other_sections_prompt))  # New line
    }

    # Combine all responses
    combined_data = {
        'personal_info': responses['personal'].get('personal_info', {}) if responses['personal'] else {},
        'professional_summary': responses['personal'].get('professional_summary', '') if responses['personal'] else '',
        'career_objective': responses['personal'].get('career_objective', '') if responses['personal'] else '',
        'education': responses['personal'].get('education', []) if responses['personal'] else [],
        'experience': responses['experience'].get('experience', []) if responses['experience'] else [],
        'projects': responses['projects'].get('projects', []) if responses['projects'] else [],
        'skills': responses['skills_cert'].get('skills', {}) if responses['skills_cert'] else {},
        'courses_and_certifications': responses['skills_cert'].get('courses_and_certifications', []) if responses[
            'skills_cert'] else [],
        'additional_sections': responses['other_sections'].get('additional_sections', {}) if responses[
            'other_sections'] else {}  # New line
    }

    # Update expected_structure
    expected_structure = {
        "personal_info": {
            "name": "Unknown",
            "email": "",
            "phone": "",
            "address": "",
            "LinkedIn": "",
            "date_of_birth": "",
            "nationality": "",
            "father_name": "",
        },
        "professional_summary": "",
        "career_objective": "",
        "education": [
            {
                "degree": "",
                "institution": "",
                "year": "",
                "details": ""
            }
        ],
        "experience": [
            {
                "title": "",
                "company": "",
                "duration": "",
                "responsibilities": []
            }
        ],
        "projects": [
            {
                "name": "",
                "description": "",
                "technologies": ""
            }
        ],
        "skills": {
            "technical": [],
            "soft": []
        },
        "courses_and_certifications": [
            {
                "name": "",
                "issuer": "",
                "year": "",
                "type": ""
            }
        ],
        # New section
        "additional_sections": {
            "achievements": [
                {
                    "title": "",
                    "description": "",
                    "year": ""
                }
            ],
            "volunteer_work": [
                {
                    "organization": "",
                    "role": "",
                    "duration": "",
                    "description": ""
                }
            ],
            "languages": [
                {
                    "language": "",
                    "proficiency": ""
                }
            ],
            "awards": [
                {
                    "name": "",
                    "issuer": "",
                    "year": ""
                }
            ],
            "publications": [
                {
                    "title": "",
                    "authors": "",
                    "publication_venue": "",
                    "year": ""
                }
            ],
            "professional_memberships": [
                {
                    "organization": "",
                    "role": "",
                    "year": ""
                }
            ],
            "interests_and_hobbies": [
                {
                    "name": "",
                    "description": ""
                }
            ]
        }
    }

    # Fill in missing subfields with defaults
    def fill_defaults(data, structure):
        if isinstance(data, list) and isinstance(structure, list):
            default_item = structure[0] if structure else {}
            return [fill_defaults(item, default_item) for item in data]
        elif isinstance(data, dict) and isinstance(structure, dict):
            return {
                key: fill_defaults(data.get(key, ""), structure[key])
                for key in structure
            }
        else:
            return data if data else ""

    filled_data = fill_defaults(combined_data, expected_structure)

    # Validate LinkedIn URL
    linkedin_url = filled_data['personal_info'].get('LinkedIn', '')
    if linkedin_url and not re.match(r'https?://(www\.)?linkedin\.com/', linkedin_url):
        filled_data['personal_info']['other_links'].append(linkedin_url)
        filled_data['personal_info']['LinkedIn'] = ""

    return filled_data


@st.cache_resource
def extract_images_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    images = []
    for page in pdf_reader.pages:
        if "/XObject" in page["/Resources"]:
            x_objects = page["/Resources"]["/XObject"].get_object()
            for obj in x_objects:
                if x_objects[obj]["/Subtype"] == "/Image":
                    try:
                        image_data = x_objects[obj].get_data()
                        image = PILImage.open(io.BytesIO(image_data))
                        images.append(image)
                    except Exception:
                        # Skip unsupported or corrupted images
                        continue
    return images


@st.cache_resource
def extract_images_from_docx(docx_file):
    doc = Document(docx_file)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image = PILImage.open(io.BytesIO(image_data))
                images.append(image)
            except Exception:
                # Skip unsupported or corrupted images
                continue
    return images


def create_word_document(data, profile_picture=None, logo_path=None):
    doc = Document()

    def add_horizontal_line(paragraph):
        """Add a horizontal line after the section headers"""
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bottom_border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '20')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        bottom_border.append(bottom)
        pPr.append(bottom_border)

    # Set up header with centered logo for all pages if logo is provided
    if logo_path:
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(2))  # Adjust size for display

    # Convert PIL Image to bytes stream if provided
    if profile_picture:
        image_stream = io.BytesIO()
        profile_picture.save(image_stream, format='JPEG')
        image_stream.seek(0)

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.30)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Create and configure styles
    styles = doc.styles

    # Base style for inheritance
    if 'Resume Base' not in styles:
        base_style = styles.add_style('Resume Base', WD_STYLE_TYPE.PARAGRAPH)
        base_style.font.name = 'Calibri'
        base_style.font.size = Pt(12)
        base_style.paragraph_format.space_after = Pt(0)  # Removed default spacing
        base_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Subheading Style (for job titles, education details, etc.)
    if 'Resume Subheading' not in styles:
        subheading_style = styles.add_style('Resume Subheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.base_style = styles['Resume Base']
        subheading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subheading_style.font.bold = True
        subheading_style.paragraph_format.space_after = Pt(3)  # Minimal space after subheading

    # Name Header Style
    if 'Resume Name' not in styles:
        name_style = styles.add_style('Resume Name', WD_STYLE_TYPE.PARAGRAPH)
        name_style.base_style = styles['Resume Base']
        name_style.font.size = Pt(24)
        name_style.font.bold = True
        name_style.paragraph_format.space_after = Pt(3)
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Contact Info Style
    if 'Resume Contact' not in styles:
        contact_style = styles.add_style('Resume Contact', WD_STYLE_TYPE.PARAGRAPH)
        contact_style.base_style = styles['Resume Base']
        contact_style.paragraph_format.space_after = Pt(6)
        contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Section Header Style
    if 'Resume Section' not in styles:
        section_style = styles.add_style('Resume Section', WD_STYLE_TYPE.PARAGRAPH)
        section_style.base_style = styles['Resume Base']
        section_style.font.size = Pt(14)
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Pt(6)
        section_style.paragraph_format.space_after = Pt(3)
        section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Content Style
    if 'Resume Content' not in styles:
        content_style = styles.add_style('Resume Content', WD_STYLE_TYPE.PARAGRAPH)
        content_style.base_style = styles['Resume Base']
        content_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        content_style.paragraph_format.space_after = Pt(3)

    # Bullet Style (with hanging indent for multi-line bullets)
    if 'Resume Bullet' not in styles:
        bullet_style = styles.add_style('Resume Bullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.base_style = styles['Resume Base']
        bullet_style.paragraph_format.left_indent = Inches(0.75)
        bullet_style.paragraph_format.first_line_indent = Inches(-0.2)
        bullet_style.paragraph_format.space_after = Pt(3)

    # Add header with photo
    if profile_picture:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.allow_autofit = True

        # Set column widths
        table.columns[0].width = Inches(5.5)
        table.columns[1].width = Inches(0.2)

        # Left cell - Name and contact info
        left_cell = table.cell(0, 0)
        name_para = left_cell.paragraphs[0]
        name_para.style = styles['Resume Name']
        name_para.add_run(data['personal_info']['name'])

        # Contact Info
        contact_para = left_cell.add_paragraph(style='Resume Contact')
        if data['personal_info']['address']:
            contact_para.add_run('Address: ' + data['personal_info']['address'] + '\n')
        if data['personal_info']['LinkedIn']:
            contact_para.add_run('LinkedIn: ' + data['personal_info']['LinkedIn'] + '\n')
        if data['personal_info']['date_of_birth']:
            contact_para.add_run('Date of Birth: ' + data['personal_info']['date_of_birth'] + '\n')
        if data['personal_info']['nationality']:
            contact_para.add_run('Nationality: ' + data['personal_info']['nationality'] + '\n')
        if data['personal_info']['father_name']:
            contact_para.add_run('Father\'s Name: ' + data['personal_info']['father_name'] + '\n')

        # Right cell - Photo
        right_cell = table.cell(0, 1)
        paragraph = right_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(1.5), height=Inches(1.5))
    else:
        name_para = doc.add_paragraph(style='Resume Name')
        name_para.add_run(data['personal_info']['name'])

        contact_para = doc.add_paragraph(style='Resume Contact')
        if data['personal_info']['address']:
            contact_para.add_run('Address ' + data['personal_info']['address'] + '\n')
        if data['personal_info']['LinkedIn']:
            contact_para.add_run('LinkedIn: ' + data['personal_info']['LinkedIn'] + '\n')
        if data['personal_info']['date_of_birth']:
            contact_para.add_run('Date of Birth: ' + data['personal_info']['date_of_birth'] + '\n')
        if data['personal_info']['nationality']:
            contact_para.add_run('Nationality: ' + data['personal_info']['nationality'] + '\n')
        if data['personal_info']['father_name']:
            contact_para.add_run('Father\'s Name: ' + data['personal_info']['father_name'] + '\n')

    # Professional Summary
    if data['professional_summary']:
        summary_header = doc.add_paragraph('PROFESSIONAL SUMMARY', style='Resume Section')
        add_horizontal_line(summary_header)
        doc.add_paragraph(data['professional_summary'], style='Resume Content')

    if data['career_objective']:
        summary_header = doc.add_paragraph('CAREER OBJECTIVE', style='Resume Section')
        add_horizontal_line(summary_header)
        doc.add_paragraph(data['career_objective'], style='Resume Content')

    # Education
    edu_header = doc.add_paragraph('EDUCATION', style='Resume Section')
    add_horizontal_line(edu_header)
    for edu in data['education']:
        p = doc.add_paragraph(style='Resume Subheading')  # Changed to left-aligned subheading
        p.add_run(f"{edu['degree']} - {edu['institution']}")

        details_para = doc.add_paragraph(style='Resume Content')
        details_run = details_para.add_run(f"{edu['year']}")
        if edu['details']:
            details_run.add_text(f" | {edu['details']}")

    # Skills
    skills_header = doc.add_paragraph('SKILLS', style='Resume Section')
    add_horizontal_line(skills_header)
    technical_skills = data['skills']['technical']
    if technical_skills:
        for category, skills in technical_skills.items():
            if skills:  # Only add categories with skills
                p = doc.add_paragraph(style='Resume Content')
                p.add_run(f"{category}: ").bold = True
                p.add_run(', '.join(skills))

    soft_skills = data['skills']['soft']
    if soft_skills:
        p = doc.add_paragraph(style='Resume Content')
        p.add_run('Soft Skills: ').bold = True
        p.add_run(', '.join(soft_skills))

    # Professional Experience
    exp_header = doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Resume Section')
    add_horizontal_line(exp_header)
    for exp in data['experience']:
        p = doc.add_paragraph(style='Resume Subheading')  # Changed to left-aligned subheading
        p.add_run(f"{exp['title']} - {exp['company']}")

        details_para = doc.add_paragraph(style='Resume Content')
        details_para.add_run(exp['duration'])

        for resp in exp['responsibilities']:
            bullet = doc.add_paragraph(style='Resume Bullet')
            bullet.add_run('• ' + resp)

    # Projects
    if data['projects']:
        prj_header = doc.add_paragraph('PROJECTS', style='Resume Section')
        add_horizontal_line(prj_header)
        for project in data['projects']:
            p = doc.add_paragraph(style='Resume Subheading')  # Changed to left-aligned subheading
            p.add_run(project['name'])

            if project['description']:
                desc_para = doc.add_paragraph(style='Resume Content')
                desc_para.add_run(project['description'])

            if project['technologies']:
                tech_para = doc.add_paragraph(style='Resume Bullet')
                tech_para.add_run('Technologies: ').bold = True
                tech_para.add_run(project['technologies'])

    # Certifications
    if data['courses_and_certifications']:
        cc_header = doc.add_paragraph('COURSES AND CERTIFICATIONS', style='Resume Section')
        add_horizontal_line(cc_header)
        for cert in data['courses_and_certifications']:
            p = doc.add_paragraph(style='Resume Bullet')
            p.add_run('• ' + f"{cert['name']} - {cert['issuer']} ({cert['year']}) - {cert['type'].capitalize()}")
    # Additional Sections
    additional_sections = data.get('additional_sections', {})

    # Achievements
    achievements = additional_sections.get('achievements', [])
    if achievements and achievements[0]['title']:
        ach_header = doc.add_paragraph('ACHIEVEMENTS', style='Resume Section')
        add_horizontal_line(ach_header)
        for achievement in achievements:
            p = doc.add_paragraph(style='Resume Bullet')
            achievement_text = achievement['title']
            if achievement['year']:
                achievement_text += f" ({achievement['year']})"
            if achievement['description']:
                achievement_text += f" - {achievement['description']}"
            p.add_run('• ' + achievement_text)

    # Volunteer Work
    volunteer_work = additional_sections.get('volunteer_work', [])
    if volunteer_work and volunteer_work[0]['organization']:
        vol_header = doc.add_paragraph('VOLUNTEER WORK', style='Resume Section')
        add_horizontal_line(vol_header)
        for volunteer in volunteer_work:
            p = doc.add_paragraph(style='Resume Content')
            p.add_run(f"{volunteer['role']} - {volunteer['organization']}").bold = True
            if volunteer['duration']:
                p.add_run(f"\n{volunteer['duration']}")
            if volunteer['description']:
                bullet = doc.add_paragraph(style='Resume Bullet')
                bullet.add_run('• ' + volunteer['description'])

    # Languages
    languages = additional_sections.get('languages', [])
    if languages and languages[0]['language']:
        lang_header = doc.add_paragraph('LANGUAGES', style='Resume Section')
        add_horizontal_line(lang_header)
        p = doc.add_paragraph(style='Resume Content')
        language_text = []
        for lang in languages:
            lang_entry = lang['language']
            if lang['proficiency']:
                lang_entry += f" ({lang['proficiency']})"
            language_text.append(lang_entry)
        p.add_run(', '.join(language_text))

    # Awards
    awards = additional_sections.get('awards', [])
    if awards and awards[0]['name']:
        award_header = doc.add_paragraph('AWARDS', style='Resume Section')
        add_horizontal_line(award_header)
        for award in awards:
            p = doc.add_paragraph(style='Resume Bullet')
            award_text = award['name']
            if award['issuer']:
                award_text += f" - {award['issuer']}"
            if award['year']:
                award_text += f" ({award['year']})"
            p.add_run('• ' + award_text)

    # Publications
    publications = additional_sections.get('publications', [])
    if publications and publications[0]['title']:
        pub_header = doc.add_paragraph('PUBLICATIONS', style='Resume Section')
        add_horizontal_line(pub_header)
        for pub in publications:
            p = doc.add_paragraph(style='Resume Bullet')
            pub_text = pub['title']
            if pub['authors']:
                pub_text += f" by {pub['authors']}"
            if pub['publication_venue']:
                pub_text += f" in {pub['publication_venue']}"
            if pub['year']:
                pub_text += f" ({pub['year']})"
            p.add_run('• ' + pub_text)

    # Professional Memberships
    memberships = additional_sections.get('professional_memberships', [])
    if memberships and memberships[0]['organization']:
        pm_header = doc.add_paragraph('PROFESSIONAL MEMBERSHIPS', style='Resume Section')
        add_horizontal_line(pm_header)
        for membership in memberships:
            p = doc.add_paragraph(style='Resume Bullet')
            membership_text = membership['organization']
            if membership['role']:
                membership_text += f" - {membership['role']}"
            if membership['year']:
                membership_text += f" ({membership['year']})"
            p.add_run('• ' + membership_text)
    interests_and_hobbies = additional_sections.get('interests_and_hobbies', [])
    if interests_and_hobbies and interests_and_hobbies[0]['name']:
        interests_header = doc.add_paragraph('INTERESTS AND HOBBIES', style='Resume Section')
        add_horizontal_line(interests_header)
        p = doc.add_paragraph(style='Resume Content')

        # If there's a description, include it in parentheses
        interest_text = []
        for interest in interests_and_hobbies:
            interest_entry = interest['name']
            if interest['description']:
                interest_entry += f" ({interest['description']})"
            interest_text.append(interest_entry)

        p.add_run(', '.join(interest_text))
    # Save document
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


def create_pdf_document(data, profile_picture=None, logo_path=None):
    # Create a BytesIO object to store the PDF
    pdf_buffer = io.BytesIO()

    # Create the PDF document with margins similar to Word document
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=letter,
        rightMargin=72,  # 1 inch
        leftMargin=72,  # 1 inch
        topMargin=65,  # 0.5 inch
        bottomMargin=72  # 1 inch
    )

    pdfmetrics.registerFont(TTFont('Calibri', 'calibri.ttf'))
    pdfmetrics.registerFont(TTFont('CalibriB', 'calibri_bold.ttf'))

    custom_styles = {
        'ResumeName': ParagraphStyle(
            name='ResumeName',
            fontName='CalibriB',
            fontSize=18,
            leading=28,
            spaceAfter=6,
            alignment=TA_LEFT
        ),
        'ResumeContact': ParagraphStyle(
            name='ResumeContact',
            fontName='Calibri',
            fontSize=12,
            leading=14,
            spaceAfter=6,
            alignment=TA_LEFT
        ),
        'ResumeSection': ParagraphStyle(
            name='ResumeSection',
            fontName='CalibriB',
            fontSize=14,
            leading=16,
            spaceBefore=6,
            spaceAfter=3,
            textColor=colors.black,
            alignment=TA_LEFT
        ),
        'ResumeContent': ParagraphStyle(
            name='ResumeContent',
            fontName='Calibri',
            fontSize=12,
            leading=14,
            spaceAfter=3,
            alignment=TA_JUSTIFY
        ),
        'ResumeBullet': ParagraphStyle(
            name='ResumeBullet',
            fontName='Calibri',
            fontSize=12,
            leading=14,
            leftIndent=36,
            firstLineIndent=-18,
            spaceAfter=3,
            alignment=TA_JUSTIFY
        ),
        'ResumeSubheading': ParagraphStyle(
            name='ResumeSubheading',
            fontName='CalibriB',
            fontSize=12,
            leading=14,
            spaceAfter=3,
            alignment=TA_LEFT
        )
    }
    # Build the document content
    story = []

    # Profile picture processing
    profile_pic_temp = None
    if profile_picture:
        try:
            # Convert PIL Image to bytes if needed
            if isinstance(profile_picture, PILImage.Image):
                img_byte_arr = io.BytesIO()
                profile_picture.save(img_byte_arr, format=profile_picture.format or 'PNG')
                profile_picture = img_byte_arr.getvalue()
        except Exception as e:
            st.error(f"Error processing profile picture: {str(e)}")
            profile_picture = None

    # Calculate available width
    available_width = letter[0] - doc.leftMargin - doc.rightMargin
    contact_width = available_width - 120 if profile_picture else available_width

    # Prepare contact information
    contact_info = []
    if data['personal_info']['address']:
        contact_info.append(f"Address: {data['personal_info']['address']}")
    if data['personal_info']['LinkedIn']:
        contact_info.append(f"LinkedIn: {data['personal_info']['LinkedIn']}")
    if data['personal_info']['date_of_birth']:
        contact_info.append(f"DOB: {data['personal_info']['date_of_birth']}")
    if data['personal_info']['nationality']:
        contact_info.append(f"Nationality: {data['personal_info']['nationality']}")
    if data['personal_info']['father_name']:
        contact_info.append(f"Father's Name: {data['personal_info']['father_name']}")

    # Create single paragraph combining name and contact info
    header_text = (
        f"<font size='24'><b>{data['personal_info']['name']}</b></font><br/>"
        f"{'<br/>'.join(contact_info)}"
    )

    header_paragraph = Paragraph(header_text, custom_styles['ResumeContent'])

    if profile_picture:
        # Create a temporary file for profile picture
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            if isinstance(profile_picture, bytes):
                tmp_file.write(profile_picture)
            else:
                PILImage.fromarray(profile_picture).save(tmp_file, format='PNG')
            tmp_file.flush()
            profile_pic_temp = tmp_file.name

            # Create image object for table
            img = Image(profile_pic_temp)
            img.drawHeight = 80
            img.drawWidth = 80

            # Table data with profile picture
            table_data = [[header_paragraph, img]]
    else:
        # Table data without profile picture
        table_data = [[header_paragraph]]

    # Create table with appropriate column widths
    col_widths = [contact_width, 100] if profile_picture else [contact_width]

    # Create and style the table
    header_table = Table(table_data, colWidths=col_widths)
    # Modify the table creation and styling
    if profile_picture:
        col_widths = [contact_width, 100]
        header_table = Table(table_data, colWidths=col_widths)
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), -5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
    else:
        # When no profile picture, use standard padding
        header_table = Table(table_data)
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))

    story.append(header_table)
    story.append(Spacer(1, 20))

    # Professional Summary
    if data['professional_summary']:
        story.append(Paragraph('PROFESSIONAL SUMMARY', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        story.append(Paragraph(data['professional_summary'], custom_styles['ResumeContent']))
    if data['career_objective']:
        story.append(Paragraph('CAREER OBJECTIVE', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        story.append(Paragraph(data['career_objective'], custom_styles['ResumeContent']))

    # Education
    story.append(Paragraph('EDUCATION', custom_styles['ResumeSection']))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
    for edu in data['education']:
        edu_text = f"{edu['degree']} - {edu['institution']}"
        story.append(Paragraph(edu_text, custom_styles['ResumeSubheading']))

        edu_details = []
        if edu['year']:
            edu_details.append(edu['year'])
        if edu['details']:
            edu_details.append(edu['details'])

        if edu_details:
            story.append(Paragraph(' | '.join(edu_details), custom_styles['ResumeContent']))

    # Skills
    story.append(Paragraph('SKILLS', custom_styles['ResumeSection']))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
    # Technical Skills
    if data['skills']['technical']:
        for category, skills in data['skills']['technical'].items():
            if skills:
                # Create a bold style for just the category
                category_bold_style = ParagraphStyle(
                    'TechnicalSkillCategoryBold',
                    parent=custom_styles['ResumeContent'],
                    fontName='CalibriB'
                )

                # Create a normal style for skills
                skills_normal_style = custom_styles['ResumeContent']

                # Create a composite paragraph with bold category and normal skills
                skills_text = f"<font name='CalibriB'>{category}:</font> {', '.join(skills)}"
                story.append(Paragraph(skills_text, skills_normal_style))
                story.append(Spacer(1, 6))
    # Soft Skills
    if data['skills']['soft']:
        soft_skills_text = f"<font name='CalibriB'>Soft Skills:</font> {', '.join(data['skills']['soft'])}"
        story.append(Paragraph(soft_skills_text, custom_styles['ResumeContent']))
        story.append(Spacer(1, 6))

    # Professional Experience
    story.append(Paragraph('PROFESSIONAL EXPERIENCE', custom_styles['ResumeSection']))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
    for exp in data['experience']:
        # Job title and company
        story.append(Paragraph(
            f"{exp['title']} - {exp['company']}",
            custom_styles['ResumeSubheading']
        ))

        if exp.get('duration'):
            story.append(Paragraph(exp['duration'], custom_styles['ResumeContent']))

        # Responsibilities
        for resp in exp['responsibilities']:
            story.append(Paragraph(f"• {resp}", custom_styles['ResumeBullet']))

    # Projects
    if data['projects']:
        story.append(Paragraph('PROJECTS', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        for project in data['projects']:
            # Project name
            story.append(Paragraph(project['name'], custom_styles['ResumeSubheading']))

            # Description
            if project['description']:
                story.append(Paragraph(project['description'], custom_styles['ResumeContent']))

            # Technologies
            if project['technologies']:
                story.append(Paragraph(
                    f"Technologies: {project['technologies']}",
                    custom_styles['ResumeContent']
                ))

            story.append(Spacer(1, 6))

    # Certifications
    if data['courses_and_certifications']:
        story.append(Paragraph('COURSES AND CERTIFICATIONS', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        for cert in data['courses_and_certifications']:
            cert_text = f"{cert['name']} - {cert['issuer']} ({cert['year']}) - {cert['type'].capitalize()}"
            story.append(Paragraph(f"• {cert_text}", custom_styles['ResumeBullet']))
        story.append(Spacer(1, 6))

    # Additional Sections
    additional_sections = data.get('additional_sections', {})

    # Achievements
    achievements = additional_sections.get('achievements', [])
    if achievements and achievements[0]['title']:
        story.append(Paragraph('ACHIEVEMENTS', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        for achievement in achievements:
            achievement_text = achievement['title']
            if achievement['year']:
                achievement_text += f" ({achievement['year']})"
            if achievement['description']:
                achievement_text += f" - {achievement['description']}"
            story.append(Paragraph(f"• {achievement_text}", custom_styles['ResumeBullet']))
        story.append(Spacer(1, 6))

    # Volunteer Work
    volunteer_work = additional_sections.get('volunteer_work', [])
    if volunteer_work and volunteer_work[0]['organization']:
        story.append(Paragraph('VOLUNTEER WORK', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        for volunteer in volunteer_work:
            story.append(Paragraph(
                f"{volunteer['role']} - {volunteer['organization']}",
                custom_styles['ResumeSubheading']
            ))

            if volunteer['duration']:
                story.append(Paragraph(volunteer['duration'], custom_styles['ResumeContent']))

            if volunteer['description']:
                story.append(Paragraph(f"• {volunteer['description']}", custom_styles['ResumeBullet']))

        story.append(Spacer(1, 6))

    # Languages
    languages = additional_sections.get('languages', [])
    if languages and languages[0]['language']:
        story.append(Paragraph('LANGUAGES', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        language_text = []
        for lang in languages:
            lang_entry = lang['language']
            if lang['proficiency']:
                lang_entry += f" ({lang['proficiency']})"
            language_text.append(lang_entry)
        story.append(Paragraph(', '.join(language_text), custom_styles['ResumeContent']))
        story.append(Spacer(1, 6))

    # Awards
    awards = additional_sections.get('awards', [])
    if awards and awards[0]['name']:
        story.append(Paragraph('AWARDS', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        for award in awards:
            award_text = award['name']
            if award['issuer']:
                award_text += f" - {award['issuer']}"
            if award['year']:
                award_text += f" ({award['year']})"
            story.append(Paragraph(f"• {award_text}", custom_styles['ResumeBullet']))
        story.append(Spacer(1, 6))

    # Publications
    publications = additional_sections.get('publications', [])
    if publications and publications[0]['title']:
        story.append(Paragraph('PUBLICATIONS', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        for pub in publications:
            pub_text = pub['title']
            if pub['authors']:
                pub_text += f" by {pub['authors']}"
            if pub['publication_venue']:
                pub_text += f" in {pub['publication_venue']}"
            if pub['year']:
                pub_text += f" ({pub['year']})"
            story.append(Paragraph(f"• {pub_text}", custom_styles['ResumeBullet']))
        story.append(Spacer(1, 6))

    # Professional Memberships
    memberships = additional_sections.get('professional_memberships', [])
    if memberships and memberships[0]['organization']:
        story.append(Paragraph('PROFESSIONAL MEMBERSHIPS', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))
        for membership in memberships:
            membership_text = membership['organization']
            if membership['role']:
                membership_text += f" - {membership['role']}"
            if membership['year']:
                membership_text += f" ({membership['year']})"
            story.append(Paragraph(f"• {membership_text}", custom_styles['ResumeBullet']))
        story.append(Spacer(1, 6))
    interests_and_hobbies = additional_sections.get('interests_and_hobbies', [])
    if interests_and_hobbies and interests_and_hobbies[0]['name']:
        story.append(Paragraph('INTERESTS AND HOBBIES', custom_styles['ResumeSection']))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.black))

        interest_text = []
        for interest in interests_and_hobbies:
            interest_entry = interest['name']
            if interest['description']:
                interest_entry += f" ({interest['description']})"
            interest_text.append(interest_entry)

        story.append(Paragraph(', '.join(interest_text), custom_styles['ResumeContent']))
        story.append(Spacer(1, 6))

    def first_page(canvas, doc):
        """Add logo to the first page."""
        canvas.saveState()
        try:
            # Add logo if provided
            if logo_path and os.path.exists(logo_path):
                # Load and resize logo
                with PILImage.open(logo_path) as logo_img:
                    logo_width = 100
                    aspect = logo_img.height / logo_img.width
                    logo_height = logo_width * aspect

                    # Position logo at top center
                    canvas.drawImage(
                        logo_path,
                        (letter[0] - logo_width) / 2,
                        letter[1] - logo_height - 36,
                        width=logo_width,
                        height=logo_height
                    )
        except Exception as e:
            st.error(f"Error in first_page: {str(e)}")
        finally:
            canvas.restoreState()

    def later_pages(canvas, doc):
        """Add logo to later pages."""
        canvas.saveState()
        try:
            if logo_path and os.path.exists(logo_path):
                with PILImage.open(logo_path) as logo_img:
                    logo_width = 100
                    aspect = logo_img.height / logo_img.width
                    logo_height = logo_width * aspect

                    canvas.drawImage(
                        logo_path,
                        (letter[0] - logo_width) / 2,
                        letter[1] - logo_height - 36,
                        width=logo_width,
                        height=logo_height
                    )
        except Exception as e:
            st.error(f"Error in later_pages: {str(e)}")
        finally:
            canvas.restoreState()

    try:
        # Build the PDF with templates
        doc.build(
            story,
            onFirstPage=first_page,
            onLaterPages=later_pages
        )

        # Get the value of the BytesIO buffer
        pdf_bytes = pdf_buffer.getvalue()
        return pdf_bytes

    except Exception as e:
        st.error(f"Error building PDF: {str(e)}")
        return None

    finally:
        # Clean up resources
        pdf_buffer.close()
        if profile_pic_temp and os.path.exists(profile_pic_temp):
            try:
                os.unlink(profile_pic_temp)
            except Exception as e:
                st.error(f"Error cleaning up temporary file: {str(e)}")


def process_document(file, logo_path, file_extension):
    with tempfile.NamedTemporaryFile(delete=False, suffix="." + file_extension) as tmp_file:
        tmp_file.write(file.getvalue())
        tmp_file_path = tmp_file.name

    try:
        if file_extension == 'pdf':
            text = get_pdf_text(tmp_file_path)
            images = extract_images_from_pdf(tmp_file_path)
        elif file_extension == 'docx':
            text = get_docx_text(tmp_file_path)
            images = extract_images_from_docx(tmp_file_path)

        if len(text) > 7000:
            extracted_info = extract_info_with_gemini(text)
        else:
            extracted_info = extract_info_with_gemini_mini(text)

        if extracted_info:
            name = extracted_info['personal_info']['name']
            clean_name = re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')

            profile_picture = images[0] if images else None

            processed_data = ProcessedData(
                extracted_info=extracted_info,
                profile_picture=profile_picture,
                word_with_pic=create_word_document(extracted_info, profile_picture, logo_path),
                pdf_with_pic=create_pdf_document(extracted_info, profile_picture, logo_path),
                word_without_pic=create_word_document(extracted_info, None, logo_path),
                pdf_without_pic=create_pdf_document(extracted_info, None, logo_path)
            )

            return processed_data

        return None
    finally:
        os.unlink(tmp_file_path)


@st.cache_data
def cached_process_document(file, logo_path, file_extension):
    return process_document(file, logo_path, file_extension)


class ProcessedData:
    def __init__(self, extracted_info=None, profile_picture=None, word_with_pic=None,
                 pdf_with_pic=None, word_without_pic=None, pdf_without_pic=None):
        self.extracted_info = extracted_info
        self.profile_picture = profile_picture
        self.word_with_pic = word_with_pic
        self.pdf_with_pic = pdf_with_pic
        self.word_without_pic = word_without_pic
        self.pdf_without_pic = pdf_without_pic


def initialize_formatting_count():
    """Initialize or load the all-time formatting count."""
    FORMATTING_FILE = 'all_time_resume_formatting.json'

    # Create file if it doesn't exist
    if not os.path.exists(FORMATTING_FILE):
        with open(FORMATTING_FILE, 'w') as f:
            json.dump({
                'total_formats': 0,
                'first_format_timestamp': datetime.now().isoformat()
            }, f)

    # Read current formatting data
    with open(FORMATTING_FILE, 'r') as f:
        formatting_data = json.load(f)

    return formatting_data['total_formats']


def increment_formatting_count():
    """Increment the all-time formatting count."""
    FORMATTING_FILE = 'all_time_resume_formatting.json'

    # Read current formatting data
    with open(FORMATTING_FILE, 'r') as f:
        formatting_data = json.load(f)

    # Increment formatting count
    formatting_data['total_formats'] += 1

    # Write updated data
    with open(FORMATTING_FILE, 'w') as f:
        json.dump(formatting_data, f)

    return formatting_data['total_formats']


def main():
    logo_path = 'logo.png'
    all_time_formats = initialize_formatting_count()
    st.set_page_config(
        page_title="Resume Formatter- CL(V12)",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded",
        menu_items={
            'Get Help': None,
            'Report a bug': None,
            'About': None
        }
    )

    # Custom CSS with fixed header spacing
    st.markdown("""
        <style>
        .stApp {
            max-width: 100%;
            margin: 0;
            padding: 0;
            background-color: #0A192F;
        }
        /* Remove all default padding from the main container */
        .main .block-container {
            padding-top: 0 !important;
            padding-bottom: 1rem !important;
            max-width: 100% !important;
        }
        /* Remove default gaps */
        .css-18e3th9 {
            padding-top: 0 !important;
            padding-bottom: 0 !important;
        }
        .css-1d391kg {
            padding-top: 0 !important;
        }
        .main-header {
            text-align: center;
            padding: 0.5rem 0;
            color: #FFFFFF;
            margin: 0;
        }
        .main-header h1 {
            color: #FFFFFF;
            margin: 0;
            padding: 0;
        }
        .main-header p {
            color: #FFFFFF;
            margin: 0.25rem 0 0 0;
        }
        .upload-section {
            background-color: transparent;
            padding: 1rem;
            border-radius: 8px;
            margin: 0.5rem 0;
        }
        /* Updated sidebar styling */
        [data-testid="stSidebar"] {
            width: 250px !important;
            background-color: #172A46 !important;
            border-right: 1px solid #234670;
        }
        [data-testid="stSidebarNav"] {
            width: 250px !important;
        }
        /* Sidebar content color */
        [data-testid="stSidebar"] .stMarkdown {
            color: #E6F1FF;
        }
        /* Sidebar header color */
        [data-testid="stSidebar"] h1, 
        [data-testid="stSidebar"] h2, 
        [data-testid="stSidebar"] h3, 
        [data-testid="stSidebar"] h4 {
            color: #64FFDA;
        }
        /* Remove white background and default header */
        [data-testid="stAppViewContainer"] {
            background-color: #0A192F;
        }
        [data-testid="stHeader"] {
            display: none;
        }
        [data-testid="stToolbar"] {
            display: none;
        }
        /* Force remove top margin from first element */
        .element-container:first-child {
            margin-top: 0 !important;
        }
        /* Additional spacing fixes */
        .stMarkdown {
            color: #FFFFFF;
            margin-top: 0 !important;
        }
        div[data-testid="stVerticalBlock"] > div:first-child {
            margin-top: 0 !important;
            padding-top: 0 !important;
        }
        /* Fix for streamlit containers */
        .stMarkdown > div:first-child > div:first-child {
            margin-top: 0 !important;
            padding-top: 0 !important;
        }
        </style>
    """, unsafe_allow_html=True)

    # Header Section - minimal spacing
    st.markdown("""
        <div class='main-header'>
        <h1>📄 Resume Formatter- CL(V12)</h1>
        </div>
    """, unsafe_allow_html=True)

    # Sidebar for settings and help
    with st.sidebar:
        # Add logo to sidebar
        st.image(logo_path, width=200)
        st.header("⚙️ Settings & Help")
        testing_mode = st.toggle("Testing Mode", value=False)
        if st.button("⚠️ Confirm Clear Session?", key="confirm_clear"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.cache_data.clear()
            st.success("Session has been cleared successfully!")
            st.rerun()
        count_placeholder = st.sidebar.empty()

        # Initial count display
        initial_count = initialize_formatting_count()
        count_placeholder.markdown(f"**Total Resumes Formatted:** {initial_count}")

        st.markdown("---")
        st.subheader("📋 Quick Guide")
        st.markdown("""
            1. Upload a resume (PDF or Word format)
            2. Add an optional profile picture
            3. Sometimes, the system may extract multiple images from the resume. If the extracted profile picture is incorrect, use the Profile Picture Management section to upload a new profile picture.
            4. Download formatted versions
            5. When uploading a new resume, please clear the session by clicking the "Clear Session" button in the section above
            6. If the output is repeating, clear the session and refresh the page
            7. You can edit the content as needed
            8. Recommended: Upload resume in PDF format and download the formatted version in Word document format
            9. Testing mode is only for developers, do not turn it on unless you are a developer \n
            Happy Hiring!!!!
        """)
        st.markdown("---")
        st.subheader("🆘 Need Help?")
        st.markdown("""
            If you encounter any issues, please [upload the resume here](https://cirruslabsio-my.sharepoint.com/:f:/g/personal/mukund_hs_cirruslabs_io/EupXiGkB1dVGn_G6QAORzYoBJ26muiIYDPmviX8KLiGdwA?e=DHQPRC) 
            for assistance.
        """)

    # Main content area
    with st.container():
        # File upload section
        st.markdown("<div class='upload-section'>", unsafe_allow_html=True)
        upload_col1, upload_col2 = st.columns([3, 1])

        with upload_col1:
            uploaded_file = st.file_uploader(
                "📎 Upload a Resume",
                type=["pdf", "docx"],
                help="Supported formats: PDF, Word Document",
                key="document_uploader"
            )

        with upload_col2:
            st.markdown("#### Recommended Format:")
            st.markdown("PDF for uploading, Word format for downloading")
            # st.markdown("- Word format for downloading")

        # Process uploaded file
        if uploaded_file:
            processing_time = None
            st.success("New session has started.")
            st.warning(
                "⚠️ Disclaimer: The formatter may occasionally make errors. Please review the content carefully before downloading the formatted resume.")

            if 'processed_data' not in st.session_state:
                with st.spinner("Processing the resume..."):
                    start_time = time.time()
                    file_extension = uploaded_file.name.split('.')[-1].lower()
                    processed_data = cached_process_document(uploaded_file, logo_path, file_extension)
                    st.session_state.processed_data = processed_data
                    end_time = time.time()
                    try:
                        # Only increment count if testing mode is off
                        if not testing_mode:
                            updated_count = increment_formatting_count()
                            count_placeholder.markdown(f"**Total Resumes Formatted:** {updated_count}")
                        else:
                            # When in testing mode, display the initial count
                            count_placeholder.markdown(f"**Total Resumes Formatted:** {initial_count} (Testing Mode)")
                    except Exception as e:
                        st.error(f"Error updating formatting count: {e}")
                        updated_count = initial_count
                processing_time = end_time - start_time

            processed_data = st.session_state.processed_data
            if processed_data and processed_data.extracted_info:
                extracted_info = processed_data.extracted_info
                if processing_time is not None:
                    st.success(f"Resume processed and formatted successfully in {processing_time:.2f} seconds!")
                else:
                    st.success("Resume processed and formatted successfully!")

                st.markdown("# 📝 Resume Content Editor")
                if 'edit_resume' not in st.session_state:
                    st.session_state.edit_resume = False

                    # Edit Resume Button
                if st.button("✏️ Edit Resume Details"):
                    st.session_state.edit_resume = True

                    # Only show editing section if edit mode is active
                if st.session_state.edit_resume:
                    #     #Close Edit Mode Button
                    #     if st.button("❌ Cancel Editing"):
                    #         st.session_state.edit_resume = False
                    #         return
                    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs([
                        "🧑 Personal Info",
                        "🖥 Professional Summary",
                        "🎯 Career Objective",
                        "🎓 Education",
                        "💼 Experience",
                        "🔃 Projects",
                        "💻 Skills",
                        "🏆 Certifications",
                        "➕ Additional Info"
                    ])

                    # Personal Information Tab (remains the same as previous implementation)
                    with tab1:
                        st.subheader("✏️ Personal Details")
                        col1, col2 = st.columns(2)
                        with col1:
                            extracted_info['personal_info']['name'] = st.text_input(
                                "Full Name", extracted_info['personal_info'].get('name', "")
                            )
                            extracted_info['personal_info']['email'] = st.text_input(
                                "Email", extracted_info['personal_info'].get('email', "")
                            )
                            extracted_info['personal_info']['address'] = st.text_input(
                                "Address", extracted_info['personal_info'].get('address', "")
                            )
                            # Adding Date of Birth input
                            extracted_info['personal_info']['date_of_birth'] = st.text_input(
                                "Date of Birth", extracted_info['personal_info'].get('date_of_birth', "")
                            )

                        with col2:
                            extracted_info['personal_info']['phone'] = st.text_input(
                                "Phone", extracted_info['personal_info'].get('phone', "")
                            )
                            extracted_info['personal_info']['LinkedIn'] = st.text_input(
                                "LinkedIn Profile", extracted_info['personal_info'].get('LinkedIn', "")
                            )
                            # Adding Nationality input
                            extracted_info['personal_info']['nationality'] = st.text_input(
                                "Nationality", extracted_info['personal_info'].get('nationality', "")
                            )
                            # Adding Father's Name input
                            extracted_info['personal_info']['father_name'] = st.text_input(
                                "Father's Name", extracted_info['personal_info'].get('father_name', "")
                            )
                    with tab2:
                        st.subheader("📋 Professional Summary")
                        extracted_info['professional_summary'] = st.text_area(
                            "Professional Summary",
                            value=extracted_info.get('professional_summary', ""),
                            height=200
                        )
                    with tab3:
                        st.subheader("🎯 Career Objective")
                        extracted_info['career_objective'] = st.text_area(
                            "Career Objective",
                            value=extracted_info.get('career_objective', ""),
                            height=150
                        )
                    # Education Tab (from previous implementation)
                    with tab4:
                        st.subheader("🎓 Educational Background")
                        education = extracted_info.get('education', [])

                        for i in range(len(education)):
                            with st.expander(f"Education {i + 1}"):
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    education[i]['degree'] = st.text_input(f"Degree", education[i].get('degree', ""),
                                                                           key=f"edu_degree_{i}")
                                with col2:
                                    education[i]['institution'] = st.text_input(f"Institution",
                                                                                education[i].get('institution', ""),
                                                                                key=f"edu_institution_{i}")
                                with col3:
                                    education[i]['year'] = st.text_input(f"Graduation Year",
                                                                         education[i].get('year', ""),
                                                                         key=f"edu_year_{i}")
                                education[i]['details'] = st.text_area(f"Additional Details",
                                                                       education[i].get('details', ""),
                                                                       key=f"edu_details_{i}")

                                if st.button(f"🗑️ Remove Education {i + 1}", key=f"delete_edu_{i}"):
                                    education.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Education"):
                            education.append({
                                'degree': '',
                                'institution': '',
                                'year': '',
                                'details': ''
                            })
                            st.rerun()

                        extracted_info['education'] = education

                    # Experience Tab (similar to previous implementation)
                    with tab5:
                        st.subheader("💼 Professional Experience")
                        experience = extracted_info.get('experience', [])

                        for i in range(len(experience)):
                            with st.expander(f"Experience {i + 1}"):
                                col1, col2 = st.columns(2)
                                with col1:
                                    experience[i]['title'] = st.text_input(f"Job Title", experience[i].get('title', ""),
                                                                           key=f"exp_title_{i}")
                                    experience[i]['company'] = st.text_input(f"Company",
                                                                             experience[i].get('company', ""),
                                                                             key=f"exp_company_{i}")
                                with col2:
                                    experience[i]['duration'] = st.text_input(f"Duration",
                                                                              experience[i].get('duration', ""),
                                                                              key=f"exp_duration_{i}")

                                responsibilities = "\n".join(experience[i].get('responsibilities', []))
                                experience[i]['responsibilities'] = st.text_area(
                                    f"Key Responsibilities", responsibilities, key=f"exp_resp_{i}"
                                ).split("\n")

                                if st.button(f"🗑️ Remove Experience {i + 1}", key=f"delete_exp_{i}"):
                                    experience.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Experience"):
                            experience.append({
                                'title': '',
                                'company': '',
                                'duration': '',
                                'responsibilities': []
                            })
                            st.rerun()

                        extracted_info['experience'] = experience

                    # Projects Tab
                    with tab6:
                        st.subheader("🚀 Projects")
                        projects = extracted_info.get('projects', [])

                        for i in range(len(projects)):
                            with st.expander(f"Project {i + 1}"):
                                projects[i]['name'] = st.text_input(f"Project Name", projects[i].get('name', ""),
                                                                    key=f"proj_name_{i}")
                                projects[i]['description'] = st.text_area(f"Description",
                                                                          projects[i].get('description', ""),
                                                                          key=f"proj_desc_{i}")
                                projects[i]['technologies'] = st.text_input(
                                    f"Technologies Used", projects[i].get('technologies', ""), key=f"proj_tech_{i}"
                                )

                                if st.button(f"🗑️ Remove Project {i + 1}", key=f"delete_proj_{i}"):
                                    projects.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Project"):
                            projects.append({
                                'name': '',
                                'description': '',
                                'technologies': ''
                            })
                            st.rerun()

                        extracted_info['projects'] = projects

                    # Skills Tab
                    with tab7:
                        st.subheader("🛠️ Skills")

                        # Get all skill categories
                        technical_skills = extracted_info['skills'].get('technical', {})

                        # Convert categories to a list for easier row management
                        categories = list(technical_skills.keys())

                        # Iterate through categories in rows of 2
                        for i in range(0, len(categories), 3):
                            # Create a row of 2 columns
                            cols = st.columns(2)

                            # Process first category in the row
                            with cols[0]:
                                category = categories[i]
                                skills = technical_skills[category]
                                skills_str = ", ".join(skills) if skills else ""

                                edited_skills_str = st.text_area(
                                    f"{category} Skills",
                                    skills_str,
                                    key=f"skills_{category}"
                                )

                                # Process edited skills
                                edited_skills = [
                                    skill.strip() for skill in edited_skills_str.split(',')
                                    if skill.strip()
                                ]

                                if edited_skills:
                                    technical_skills[category] = edited_skills

                            # Process second category in the row (if exists)
                            if i + 1 < len(categories):
                                with cols[1]:
                                    category = categories[i + 1]
                                    skills = technical_skills[category]
                                    skills_str = ", ".join(skills) if skills else ""

                                    edited_skills_str = st.text_area(
                                        f"{category} Skills",
                                        skills_str,
                                        key=f"skills_{category}"
                                    )

                                    # Process edited skills
                                    edited_skills = [
                                        skill.strip() for skill in edited_skills_str.split(',')
                                        if skill.strip()
                                    ]

                                    if edited_skills:
                                        technical_skills[category] = edited_skills

                        # Update the skills in extracted_info
                        extracted_info['skills']['technical'] = technical_skills

                        # Soft Skills (Optional: you can adjust placement as needed)
                        st.markdown("#### Soft Skills")
                        soft_skills_str = ", ".join(extracted_info['skills'].get('soft', []))

                        edited_soft_skills_str = st.text_area(
                            "Soft Skills",
                            soft_skills_str,
                            key="soft_skills"
                        )

                        extracted_info['skills']['soft'] = [
                            skill.strip() for skill in edited_soft_skills_str.split(',')
                            if skill.strip()
                        ]
                    # Certifications Tab
                    with tab8:
                        st.subheader("🏆 Certifications")
                        certifications = extracted_info.get('courses_and_certifications', [])

                        for i in range(len(certifications)):
                            with st.expander(f"Certification {i + 1}"):
                                col1, col2 = st.columns(2)
                                with col1:
                                    certifications[i]['name'] = st.text_input(f"Certification Name",
                                                                              certifications[i].get('name', ""),
                                                                              key=f"cert_name_{i}")
                                    certifications[i]['issuer'] = st.text_input(f"Issuer",
                                                                                certifications[i].get('issuer', ""),
                                                                                key=f"cert_issuer_{i}")
                                with col2:
                                    certifications[i]['year'] = st.text_input(f"Year",
                                                                              certifications[i].get('year', ""),
                                                                              key=f"cert_year_{i}")
                                    certifications[i]['type'] = st.text_input(f"Type",
                                                                              certifications[i].get('type', ""),
                                                                              key=f"cert_type_{i}")

                                if st.button(f"🗑️ Remove Certification {i + 1}", key=f"delete_cert_{i}"):
                                    certifications.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Certification"):
                            certifications.append({
                                'name': '',
                                'issuer': '',
                                'year': '',
                                'type': ''
                            })
                            st.rerun()

                        extracted_info['courses_and_certifications'] = certifications

                    # Additional Information Tab
                    with tab9:
                        st.subheader("➕ Additional Information")
                        additional_sections = extracted_info.get('additional_sections', {})

                        # Achievements
                        st.markdown("#### 🏅 Achievements")
                        achievements = additional_sections.get('achievements', [])

                        for i in range(len(achievements)):
                            with st.expander(f"Achievement {i + 1}"):
                                achievements[i]['title'] = st.text_input(f"Title", achievements[i].get('title', ""),
                                                                         key=f"ach_title_{i}")
                                achievements[i]['description'] = st.text_area(f"Description",
                                                                              achievements[i].get('description', ""),
                                                                              key=f"ach_desc_{i}")
                                achievements[i]['year'] = st.text_input(f"Year", achievements[i].get('year', ""),
                                                                        key=f"ach_year_{i}")

                                if st.button(f"🗑️ Remove Achievement {i + 1}", key=f"delete_ach_{i}"):
                                    achievements.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Achievement"):
                            achievements.append({
                                'title': '',
                                'description': '',
                                'year': ''
                            })
                            st.rerun()

                        additional_sections['achievements'] = achievements

                        # Volunteer Work
                        st.markdown("#### 🤝 Volunteer Work")
                        volunteer_work = additional_sections.get('volunteer_work', [])

                        for i in range(len(volunteer_work)):
                            with st.expander(f"Volunteer Work {i + 1}"):
                                volunteer_work[i]['organization'] = st.text_input(f"Organization",
                                                                                  volunteer_work[i].get('organization',
                                                                                                        ""),
                                                                                  key=f"vol_org_{i}")
                                volunteer_work[i]['role'] = st.text_input(f"Role", volunteer_work[i].get('role', ""),
                                                                          key=f"vol_role_{i}")
                                volunteer_work[i]['duration'] = st.text_input(f"Duration",
                                                                              volunteer_work[i].get('duration', ""),
                                                                              key=f"vol_duration_{i}")
                                volunteer_work[i]['description'] = st.text_area(f"Description",
                                                                                volunteer_work[i].get('description',
                                                                                                      ""),
                                                                                key=f"vol_desc_{i}")

                                if st.button(f"🗑️ Remove Volunteer Work {i + 1}", key=f"delete_vol_{i}"):
                                    volunteer_work.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Volunteer Work"):
                            volunteer_work.append({
                                'organization': '',
                                'role': '',
                                'duration': '',
                                'description': ''
                            })
                            st.rerun()

                        additional_sections['volunteer_work'] = volunteer_work
                        st.markdown("#### 🌐 Languages")
                        languages = additional_sections.get('languages', [])

                        for i in range(len(languages)):
                            with st.expander(f"Language {i + 1}"):
                                languages[i]['language'] = st.text_input(f"Language", languages[i].get('language', ""),
                                                                         key=f"lang_{i}")
                                languages[i]['proficiency'] = st.text_input(f"Proficiency",
                                                                            languages[i].get('proficiency', ""),
                                                                            key=f"lang_prof_{i}")

                                if st.button(f"🗑️ Remove Language {i + 1}", key=f"delete_lang_{i}"):
                                    languages.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Language"):
                            languages.append({
                                'language': '',
                                'proficiency': ''
                            })
                            st.rerun()

                        additional_sections['languages'] = languages

                        # Awards
                        st.markdown("#### 🏆 Awards")
                        awards = additional_sections.get('awards', [])

                        for i in range(len(awards)):
                            with st.expander(f"Award {i + 1}"):
                                awards[i]['name'] = st.text_input(f"Award Name", awards[i].get('name', ""),
                                                                  key=f"award_name_{i}")
                                awards[i]['issuer'] = st.text_input(f"Issuer", awards[i].get('issuer', ""),
                                                                    key=f"award_issuer_{i}")
                                awards[i]['year'] = st.text_input(f"Year", awards[i].get('year', ""),
                                                                  key=f"award_year_{i}")

                                if st.button(f"🗑️ Remove Award {i + 1}", key=f"delete_award_{i}"):
                                    awards.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Award"):
                            awards.append({
                                'name': '',
                                'issuer': '',
                                'year': ''
                            })
                            st.rerun()

                        additional_sections['awards'] = awards

                        # Publications
                        st.markdown("#### 📚 Publications")
                        publications = additional_sections.get('publications', [])

                        for i in range(len(publications)):
                            with st.expander(f"Publication {i + 1}"):
                                publications[i]['title'] = st.text_input(f"Title", publications[i].get('title', ""),
                                                                         key=f"pub_title_{i}")
                                publications[i]['authors'] = st.text_input(f"Authors",
                                                                           publications[i].get('authors', ""),
                                                                           key=f"pub_authors_{i}")
                                publications[i]['publication_venue'] = st.text_input(f"Publication Venue",
                                                                                     publications[i].get(
                                                                                         'publication_venue', ""),
                                                                                     key=f"pub_venue_{i}")
                                publications[i]['year'] = st.text_input(f"Year", publications[i].get('year', ""),
                                                                        key=f"pub_year_{i}")

                                if st.button(f"🗑️ Remove Publication {i + 1}", key=f"delete_pub_{i}"):
                                    publications.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Publication"):
                            publications.append({
                                'title': '',
                                'authors': '',
                                'publication_venue': '',
                                'year': ''
                            })
                            st.rerun()

                        additional_sections['publications'] = publications

                        # Professional Memberships
                        st.markdown("#### 🤝 Professional Memberships")
                        professional_memberships = additional_sections.get('professional_memberships', [])

                        for i in range(len(professional_memberships)):
                            with st.expander(f"Membership {i + 1}"):
                                professional_memberships[i]['organization'] = st.text_input(f"Organization",
                                                                                            professional_memberships[
                                                                                                i].get(
                                                                                                'organization', ""),
                                                                                            key=f"membership_org_{i}")
                                professional_memberships[i]['role'] = st.text_input(f"Role",
                                                                                    professional_memberships[i].get(
                                                                                        'role',
                                                                                        ""),
                                                                                    key=f"membership_role_{i}")
                                professional_memberships[i]['year'] = st.text_input(f"Year",
                                                                                    professional_memberships[i].get(
                                                                                        'year',
                                                                                        ""),
                                                                                    key=f"membership_year_{i}")

                                if st.button(f"🗑️ Remove Membership {i + 1}", key=f"delete_membership_{i}"):
                                    professional_memberships.pop(i)
                                    st.rerun()

                        if st.button("➕ Add Professional Membership"):
                            professional_memberships.append({
                                'organization': '',
                                'role': '',
                                'year': ''
                            })
                            st.rerun()

                        additional_sections['professional_memberships'] = professional_memberships

                        # Update additional sections in extracted info
                        extracted_info['additional_sections'] = additional_sections
                        # Interests and Hobbies
                        st.markdown("#### 🌟 Interests and Hobbies")
                        interests_and_hobbies = additional_sections.get('interests_and_hobbies', [])

                        # Create a container to manage interests and hobbies
                        for i in range(len(interests_and_hobbies)):
                            with st.expander(f"Interest/Hobby {i + 1}"):
                                # Name input with dynamic key
                                interests_and_hobbies[i]['name'] = st.text_input(
                                    "Name",
                                    interests_and_hobbies[i].get('name', ""),
                                    key=f"interest_hobby_name_{i}"
                                )

                                # Description input with dynamic key
                                interests_and_hobbies[i]['description'] = st.text_input(
                                    "Description",
                                    interests_and_hobbies[i].get('description', ""),
                                    key=f"interest_hobby_description_{i}"
                                )

                                # Remove interest/hobby button
                                if st.button(f"🗑️ Remove Interest/Hobby {i + 1}", key=f"delete_interest_hobby_{i}"):
                                    interests_and_hobbies.pop(i)
                                    st.rerun()

                        # Add new interest/hobby button
                        if st.button("➕ Add Interest or Hobby"):
                            interests_and_hobbies.append({
                                'name': '',
                                'description': ''
                            })
                            st.rerun()

                        # Update additional sections with modified interests and hobbies
                        additional_sections['interests_and_hobbies'] = interests_and_hobbies
                        extracted_info['additional_sections'] = additional_sections

                    # Save Changes Button
                    if st.button("💾 Save Changes and Regenerate", type="primary"):
                        st.session_state.processed_data.extracted_info = extracted_info
                        st.success("✅ Changes saved successfully!")

                        # Regenerate Documents logic
                        with st.spinner("Regenerating documents..."):
                            profile_picture = processed_data.profile_picture
                            processed_data.word_with_pic = create_word_document(extracted_info, profile_picture,
                                                                                logo_path)
                            processed_data.pdf_with_pic = create_pdf_document(extracted_info, profile_picture,
                                                                              logo_path)
                            processed_data.word_without_pic = create_word_document(extracted_info, None, logo_path)
                            processed_data.pdf_without_pic = create_pdf_document(extracted_info, None, logo_path)
                            st.session_state.processed_data = processed_data

                # Profile Picture Management
                st.markdown("### 📸 Profile Picture Management")
                upload_col1, _ = st.columns([3, 1])

                with upload_col1:
                    # Show current profile picture status
                    has_profile_picture = hasattr(processed_data,
                                                  'profile_picture') and processed_data.profile_picture is not None

                    if has_profile_picture:
                        st.success("✅ Profile picture added successfully")
                        if st.button("❌ Remove Profile Picture"):
                            processed_data.profile_picture = None
                            processed_data.word_with_pic = processed_data.word_without_pic
                            processed_data.pdf_with_pic = processed_data.pdf_without_pic
                            st.session_state.processed_data = processed_data
                            st.success("Profile picture removed!")
                            st.rerun()

                    # Profile picture upload
                    uploaded_profile = st.file_uploader(
                        "Upload Profile Picture",
                        type=["png", "jpg", "jpeg"],
                        key="profile_uploader",
                        help="Recommended size: 200x200 pixels"
                    )

                    if uploaded_profile:
                        try:
                            profile_image = PILImage.open(uploaded_profile)
                            st.image(profile_image, caption="Preview", width=200)

                            if st.button("✅ Confirm Picture"):
                                with st.spinner("Updating documents..."):
                                    max_size = (400, 400)
                                    if profile_image.size[0] > max_size[0] or profile_image.size[1] > max_size[1]:
                                        profile_image.thumbnail(max_size, PILImage.LANCZOS)

                                    word_with_pic = create_word_document(extracted_info, profile_image, logo_path)
                                    pdf_with_pic = create_pdf_document(extracted_info, profile_image, logo_path)
                                    processed_data.word_with_pic = word_with_pic
                                    processed_data.pdf_with_pic = pdf_with_pic
                                    processed_data.profile_picture = profile_image
                                    st.session_state.processed_data = processed_data
                                    st.success("✨ Profile picture updated successfully!")
                                    st.rerun()
                        except Exception as e:
                            st.error(f"⚠️ Error processing image: {str(e)}")

                # Download Section
                st.markdown("### 📥 Download Options")
                name = extracted_info['personal_info']['name']
                clean_name = re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')

                # Preview and Download columns
                preview_col, download_col = st.columns([4, 1])

                with preview_col:
                    st.subheader("📄 Preview")
                    if processed_data.pdf_with_pic:
                        base64_pdf = base64.b64encode(processed_data.pdf_with_pic).decode('utf-8')
                        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800px" type="application/pdf"></iframe>'
                        st.markdown(pdf_display, unsafe_allow_html=True)

                with download_col:
                    st.subheader("💾 Download Files")

                    if processed_data.pdf_with_pic and processed_data.profile_picture is not None:
                        st.markdown("##### With Profile Picture")
                        st.download_button(
                            label="📑 Download PDF (with picture)",
                            data=processed_data.pdf_with_pic,
                            file_name=f"{clean_name}_resume_with_pic.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                        st.download_button(
                            label="📝 Download Word (with picture)",
                            data=processed_data.word_with_pic,
                            file_name=f"{clean_name}_resume_with_pic.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                    st.markdown("##### Without Profile Picture")
                    st.download_button(
                        label="📑 Download PDF (no picture)",
                        data=processed_data.pdf_without_pic,
                        file_name=f"{clean_name}_resume_without_pic.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    st.download_button(
                        label="📝 Download Word (no picture)",
                        data=processed_data.word_without_pic,
                        file_name=f"{clean_name}_resume_without_pic.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.error("⚠️ Failed to process the document. Please try again.")


if __name__ == "__main__":
    main()