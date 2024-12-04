from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
import io
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_word_document(data, profile_picture=None, logo_path=None):
    doc = Document()

    def add_horizontal_line(paragraph):
        """Add a horizontal line after the section headers"""
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bottom_border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '20')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        bottom_border.append(bottom)
        pPr.append(bottom_border)

    # Set up header with centered logo for all pages if logo is provided
    if logo_path:
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(2))  # Adjust size for display

    # Convert PIL Image to bytes stream if provided
    if profile_picture:
        image_stream = io.BytesIO()
        profile_picture.save(image_stream, format='JPEG')
        image_stream.seek(0)

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.30)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Create and configure styles
    styles = doc.styles

    # Base style for inheritance
    if 'Resume Base' not in styles:
        base_style = styles.add_style('Resume Base', WD_STYLE_TYPE.PARAGRAPH)
        base_style.font.name = 'Calibri'
        base_style.font.size = Pt(12)
        base_style.paragraph_format.space_after = Pt(0)  # Removed default spacing
        base_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Subheading Style (for job titles, education details, etc.)
    if 'Resume Subheading' not in styles:
        subheading_style = styles.add_style('Resume Subheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.base_style = styles['Resume Base']
        subheading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subheading_style.font.bold = True
        subheading_style.paragraph_format.space_after = Pt(3)  # Minimal space after subheading

    # Name Header Style
    if 'Resume Name' not in styles:
        name_style = styles.add_style('Resume Name', WD_STYLE_TYPE.PARAGRAPH)
        name_style.base_style = styles['Resume Base']
        name_style.font.size = Pt(24)
        name_style.font.bold = True
        name_style.paragraph_format.space_after = Pt(3)
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Contact Info Style
    if 'Resume Contact' not in styles:
        contact_style = styles.add_style('Resume Contact', WD_STYLE_TYPE.PARAGRAPH)
        contact_style.base_style = styles['Resume Base']
        contact_style.paragraph_format.space_after = Pt(6)
        contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Section Header Style
    if 'Resume Section' not in styles:
        section_style = styles.add_style('Resume Section', WD_STYLE_TYPE.PARAGRAPH)
        section_style.base_style = styles['Resume Base']
        section_style.font.size = Pt(14)
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Pt(6)
        section_style.paragraph_format.space_after = Pt(3)
        section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Content Style
    if 'Resume Content' not in styles:
        content_style = styles.add_style('Resume Content', WD_STYLE_TYPE.PARAGRAPH)
        content_style.base_style = styles['Resume Base']
        content_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        content_style.paragraph_format.space_after = Pt(3)

    # Bullet Style (with hanging indent for multi-line bullets)
    if 'Resume Bullet' not in styles:
        bullet_style = styles.add_style('Resume Bullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.base_style = styles['Resume Base']
        bullet_style.paragraph_format.left_indent = Inches(0.75)
        bullet_style.paragraph_format.first_line_indent = Inches(-0.2)
        bullet_style.paragraph_format.space_after = Pt(3)

    # Add header with photo
    if profile_picture:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.allow_autofit = True

        # Set column widths
        table.columns[0].width = Inches(5.5)
        table.columns[1].width = Inches(0.2)

        # Left cell - Name and contact info
        left_cell = table.cell(0, 0)
        name_para = left_cell.paragraphs[0]
        name_para.style = styles['Resume Name']
        name_para.add_run(data['personal_info']['name'])

        # Contact Info
        contact_para = left_cell.add_paragraph(style='Resume Contact')
        if data['personal_info']['address']:
            contact_para.add_run('Address: ' + data['personal_info']['address'] + '\n')
        if data['personal_info']['LinkedIn']:
            contact_para.add_run('LinkedIn: ' + data['personal_info']['LinkedIn'] + '\n')
        if data['personal_info']['date_of_birth']:
            contact_para.add_run('Date of Birth: ' + data['personal_info']['date_of_birth'] + '\n')
        if data['personal_info']['nationality']:
            contact_para.add_run('Nationality: ' + data['personal_info']['nationality'] + '\n')
        if data['personal_info']['father_name']:
            contact_para.add_run('Father\'s Name: ' + data['personal_info']['father_name'] + '\n')

        # Right cell - Photo
        right_cell = table.cell(0, 1)
        paragraph = right_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(1.5), height=Inches(1.5))
    else:
        name_para = doc.add_paragraph(style='Resume Name')
        name_para.add_run(data['personal_info']['name'])

        contact_para = doc.add_paragraph(style='Resume Contact')
        if data['personal_info']['address']:
            contact_para.add_run('Address ' + data['personal_info']['address'] + '\n')
        if data['personal_info']['LinkedIn']:
            contact_para.add_run('LinkedIn: ' + data['personal_info']['LinkedIn'] + '\n')
        if data['personal_info']['date_of_birth']:
            contact_para.add_run('Date of Birth: ' + data['personal_info']['date_of_birth'] + '\n')
        if data['personal_info']['nationality']:
            contact_para.add_run('Nationality: ' + data['personal_info']['nationality'] + '\n')
        if data['personal_info']['father_name']:
            contact_para.add_run('Father\'s Name: ' + data['personal_info']['father_name'] + '\n')

    # Professional Summary
    if data['professional_summary']:
        summary_header = doc.add_paragraph('PROFESSIONAL SUMMARY', style='Resume Section')
        add_horizontal_line(summary_header)
        doc.add_paragraph(data['professional_summary'], style='Resume Content')

    if data['career_objective']:
        summary_header = doc.add_paragraph('CAREER OBJECTIVE', style='Resume Section')
        add_horizontal_line(summary_header)
        doc.add_paragraph(data['career_objective'], style='Resume Content')

    # Education
    edu_header = doc.add_paragraph('EDUCATION', style='Resume Section')
    add_horizontal_line(edu_header)
    for edu in data['education']:
        p = doc.add_paragraph(style='Resume Subheading')  # Changed to left-aligned subheading
        p.add_run(f"{edu['degree']} - {edu['institution']}")

        details_para = doc.add_paragraph(style='Resume Content')
        details_run = details_para.add_run(f"{edu['year']}")
        if edu['details']:
            details_run.add_text(f" | {edu['details']}")

    # Skills
    skills_header = doc.add_paragraph('SKILLS', style='Resume Section')
    add_horizontal_line(skills_header)
    technical_skills = data['skills']['technical']
    if technical_skills:
        for category, skills in technical_skills.items():
            if skills:  # Only add categories with skills
                p = doc.add_paragraph(style='Resume Content')
                p.add_run(f"{category}: ").bold = True
                p.add_run(', '.join(skills))

    soft_skills = data['skills']['soft']
    if soft_skills:
        p = doc.add_paragraph(style='Resume Content')
        p.add_run('Soft Skills: ').bold = True
        p.add_run(', '.join(soft_skills))

    # Professional Experience
    exp_header = doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Resume Section')
    add_horizontal_line(exp_header)
    for exp in data['experience']:
        p = doc.add_paragraph(style='Resume Subheading')  # Changed to left-aligned subheading
        p.add_run(f"{exp['title']} - {exp['company']}")

        details_para = doc.add_paragraph(style='Resume Content')
        details_para.add_run(exp['duration'])

        for resp in exp['responsibilities']:
            bullet = doc.add_paragraph(style='Resume Bullet')
            bullet.add_run('• ' + resp)

    # Projects
    if data['projects']:
        prj_header = doc.add_paragraph('PROJECTS', style='Resume Section')
        add_horizontal_line(prj_header)
        for project in data['projects']:
            p = doc.add_paragraph(style='Resume Subheading')  # Changed to left-aligned subheading
            p.add_run(project['name'])

            if project['description']:
                desc_para = doc.add_paragraph(style='Resume Content')
                desc_para.add_run(project['description'])

            if project['technologies']:
                tech_para = doc.add_paragraph(style='Resume Bullet')
                tech_para.add_run('Technologies: ').bold = True
                tech_para.add_run(project['technologies'])

    # Certifications
    if data['courses_and_certifications']:
        cc_header = doc.add_paragraph('COURSES AND CERTIFICATIONS', style='Resume Section')
        add_horizontal_line(cc_header)
        for cert in data['courses_and_certifications']:
            p = doc.add_paragraph(style='Resume Bullet')
            p.add_run('• ' + f"{cert['name']} - {cert['issuer']} ({cert['year']}) - {cert['type'].capitalize()}")
    # Additional Sections
    additional_sections = data.get('additional_sections', {})

    # Achievements
    achievements = additional_sections.get('achievements', [])
    if achievements and achievements[0]['title']:
        ach_header = doc.add_paragraph('ACHIEVEMENTS', style='Resume Section')
        add_horizontal_line(ach_header)
        for achievement in achievements:
            p = doc.add_paragraph(style='Resume Bullet')
            achievement_text = achievement['title']
            if achievement['year']:
                achievement_text += f" ({achievement['year']})"
            if achievement['description']:
                achievement_text += f" - {achievement['description']}"
            p.add_run('• ' + achievement_text)

    # Volunteer Work
    volunteer_work = additional_sections.get('volunteer_work', [])
    if volunteer_work and volunteer_work[0]['organization']:
        vol_header = doc.add_paragraph('VOLUNTEER WORK', style='Resume Section')
        add_horizontal_line(vol_header)
        for volunteer in volunteer_work:
            p = doc.add_paragraph(style='Resume Content')
            p.add_run(f"{volunteer['role']} - {volunteer['organization']}").bold = True
            if volunteer['duration']:
                p.add_run(f"\n{volunteer['duration']}")
            if volunteer['description']:
                bullet = doc.add_paragraph(style='Resume Bullet')
                bullet.add_run('• ' + volunteer['description'])

    # Languages
    languages = additional_sections.get('languages', [])
    if languages and languages[0]['language']:
        lang_header = doc.add_paragraph('LANGUAGES', style='Resume Section')
        add_horizontal_line(lang_header)
        p = doc.add_paragraph(style='Resume Content')
        language_text = []
        for lang in languages:
            lang_entry = lang['language']
            if lang['proficiency']:
                lang_entry += f" ({lang['proficiency']})"
            language_text.append(lang_entry)
        p.add_run(', '.join(language_text))

    # Awards
    awards = additional_sections.get('awards', [])
    if awards and awards[0]['name']:
        award_header = doc.add_paragraph('AWARDS', style='Resume Section')
        add_horizontal_line(award_header)
        for award in awards:
            p = doc.add_paragraph(style='Resume Bullet')
            award_text = award['name']
            if award['issuer']:
                award_text += f" - {award['issuer']}"
            if award['year']:
                award_text += f" ({award['year']})"
            p.add_run('• ' + award_text)

    # Publications
    publications = additional_sections.get('publications', [])
    if publications and publications[0]['title']:
        pub_header = doc.add_paragraph('PUBLICATIONS', style='Resume Section')
        add_horizontal_line(pub_header)
        for pub in publications:
            p = doc.add_paragraph(style='Resume Bullet')
            pub_text = pub['title']
            if pub['authors']:
                pub_text += f" by {pub['authors']}"
            if pub['publication_venue']:
                pub_text += f" in {pub['publication_venue']}"
            if pub['year']:
                pub_text += f" ({pub['year']})"
            p.add_run('• ' + pub_text)

    # Professional Memberships
    memberships = additional_sections.get('professional_memberships', [])
    if memberships and memberships[0]['organization']:
        pm_header = doc.add_paragraph('PROFESSIONAL MEMBERSHIPS', style='Resume Section')
        add_horizontal_line(pm_header)
        for membership in memberships:
            p = doc.add_paragraph(style='Resume Bullet')
            membership_text = membership['organization']
            if membership['role']:
                membership_text += f" - {membership['role']}"
            if membership['year']:
                membership_text += f" ({membership['year']})"
            p.add_run('• ' + membership_text)
    interests_and_hobbies = additional_sections.get('interests_and_hobbies', [])
    if interests_and_hobbies and interests_and_hobbies[0]['name']:
        interests_header = doc.add_paragraph('INTERESTS AND HOBBIES', style='Resume Section')
        add_horizontal_line(interests_header)
        p = doc.add_paragraph(style='Resume Content')

        # If there's a description, include it in parentheses
        interest_text = []
        for interest in interests_and_hobbies:
            interest_entry = interest['name']
            if interest['description']:
                interest_entry += f" ({interest['description']})"
            interest_text.append(interest_entry)

        p.add_run(', '.join(interest_text))
    # Save document
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes